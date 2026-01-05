import docx
import re
import sys

def extract_placeholders(file_path):
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Error reading file: {e}")
        return

    placeholders = []

    # 定义匹配模式：匹配含有 * 或 X 或 下划线 的文本片段
    # 这里的正则主要为了捕获上下文，方便用户确认是哪个字段
    # 简单的逻辑：如果一行包含 * 或 X，或者连续的下划线，就认为这行包含待填信息
    
    # 扫描段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        # 查找包含 * 或 X 或 连续3个以上下划线的情况
        if '*' in text or 'X' in text or 'x' in text or '×' in text or '____' in text:
            placeholders.append(f"[段落] {text}")

    # 扫描表格
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                if '*' in text or 'X' in text or 'x' in text or '×' in text:
                    # 尝试获取表头（假设第一行是表头，或者左边第一列是标签）
                    label = "表格数据"
                    if c_idx > 0:
                        # 尝试取左边一格作为标签
                        prev_text = row.cells[c_idx-1].text.strip()
                        if prev_text:
                            label = prev_text
                    
                    placeholders.append(f"[表格] {label}: {text}")

    return placeholders

file_path = '/Users/crj/Documents/code/baogao_code/工程造价咨询报告书（结算审核）-范本.docx'
results = extract_placeholders(file_path)

print("### 提取到的待填内容预览 ###")
for item in results:
    print(item)
