from docx import Document
import re

def update_template(doc_path, output_path):
    doc = Document(doc_path)
    
    # 定义替换规则 (上下文关键字 -> 新的 Jinja2 标签)
    # 使用列表顺序执行，避免重复替换
    replacements = [
        ("本工程采用****方式招标", "本工程采用{{ bidding_method }}方式招标"),
        ("招标控制价****元", "招标控制价{{ bidding_price_control }}元"),
        ("中标价****", "中标价{{ bidding_price_winning }}"),
        ("工程质量情况：****合同约定", "工程质量情况：{{ quality_status }}合同约定"),
        ("扣除追加审计费****元", "扣除追加审计费{{ audit_fee_deduction }}元"),
        ("支付追加审计费****元", "支付追加审计费{{ audit_fee_deduction }}元"),
        ("追加审计费:****", "追加审计费:"), # 这里可能是个多余的占位，直接去掉或者替换为空
        ("咨询报告书编号：****", "咨询报告书编号：{{ report_code }}"),
        ("咨询项目委托方全称： ****", "咨询项目委托方全称： {{ client_name }}"),
        ("咨询作业期：****—****", "咨询作业期：{{ audit_period_start }}—{{ audit_period_end }}")
    ]
    
    def apply_replacements(text):
        for old, new in replacements:
            if old in text:
                print(f"Replacing: {old} -> {new}")
                text = text.replace(old, new)
        return text

    # 1. 遍历段落
    for para in doc.paragraphs:
        para.text = apply_replacements(para.text)
            
    # 2. 遍历表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.text = apply_replacements(para.text)

    doc.save(output_path)
    print(f"Updated template saved to {output_path}")

if __name__ == "__main__":
    update_template('report_template.docx', 'report_template.docx')
