from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Helper function to set cell borders.
    Usage: set_cell_border(cell, bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge, props in kwargs.items():
        tag = 'w:{}'.format(edge)
        element = tcPr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcPr.append(element)
        
        for key, value in props.items():
            element.set(qn('w:{}'.format(key)), str(value))

def create_demo_cover():
    doc = Document()
    
    # 1. 设置页面边距 (适度调整)
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 2. 标题: 工程造价咨询报告书
    # 使用空的段落占位，把标题往下推一点
    doc.add_paragraph("\n" * 4)
    
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run("工程造价咨询报告书")
    run.font.name = 'SimHei' # 黑体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(36) # 一号/小初左右
    run.bold = True

    doc.add_paragraph("\n" * 6) # 间距

    # 3. 核心信息表格 (2列)
    # 咨询项目全称, 咨询业务类别, 咨询报告日期
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    
    # 设置列宽
    # 总宽约 16cm. 左列 5cm, 右列 11cm
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)

    data = [
        ("咨询项目全称：", "{{ project_name }}"),
        ("咨询业务类别：", "工程结算审核"),
        ("咨询报告日期：", "{{ report_date }}")
    ]

    for i, (label_text, value_text) in enumerate(data):
        row = table.rows[i]
        
        # 左列：标签
        cell_label = row.cells[0]
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT # 右对齐，紧贴填空处
        run = p.add_run(label_text)
        run.font.name = 'FangSong' # 仿宋
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
        run.font.size = Pt(16) # 三号
        
        # 右列：内容
        cell_value = row.cells[1]
        p = cell_value.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中或左对齐
        run = p.add_run(value_text)
        run.font.name = 'FangSong'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
        run.font.size = Pt(16)
        
        # 给右列添加下划线 (通过设置单元格下边框实现)
        set_cell_border(cell_value, bottom={"sz": "6", "val": "single", "color": "000000", "space": "0"})
        
        # 增加行高
        row.height = Cm(1.5)

    doc.add_paragraph("\n" * 5)

    # 4. 底部 Logo 和 地址
    # Logo
    try:
        import os
        logo_path = os.path.abspath('extracted_media/word/media/image1.jpeg')
        if os.path.exists(logo_path):
            # 创建一个居中段落放图片
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 根据图片比例 (3000x428) 调整尺寸
            p_logo.add_run().add_picture(logo_path, width=Cm(14)) 
        else:
             print(f"Error: Image not found at {logo_path}")
    except Exception as e:
        print(f"Warning: Could not add logo. {repr(e)}")

    # 公司名称 (图片里可能有，如果图片包含文字则不需要再写，这里假设图片包含了 '杭州三才...')
    # 仅添加地址
    p_addr = doc.add_paragraph()
    p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_addr.add_run("地址：杭州市滨江区江汉路1786号钱龙大厦5楼 电话:0571-86603859")
    run.font.name = 'FangSong'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    run.font.size = Pt(12)

    # 保存
    output_path = 'fixed_cover_demo.docx'
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    create_demo_cover()
