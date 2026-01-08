#!/usr/bin/env python3
"""
修复模板中调整项的标点符号：最后一条用句号，其他条用分号
"""

from docx import Document

def fix_template_punctuation():
    """修改模板中的标点符号逻辑"""
    template_path = "/Users/crj/Documents/code/baogao_code/report_template.docx"
    backup_path = "/Users/crj/Documents/code/baogao_code/report_template_punctuation_backup.docx"

    print("🔍 正在读取模板文件...")
    doc = Document(template_path)

    # 备份原模板
    doc.save(backup_path)
    print(f"✅ 已备份原模板到: {backup_path}")

    # 查找包含 'adjustments' 和 'loop.index' 的段落
    target_para = None
    for i, para in enumerate(doc.paragraphs):
        if 'adjustments' in para.text and 'loop.index' in para.text:
            target_para = para
            para_index = i
            print(f"✅ 找到目标段落（第{i+1}段）:")
            print(f"   原内容: {para.text}")

            # 修改文本：替换分号为条件判断
            old_text = para.text

            # 方法：替换末尾的分号为条件判断
            # 原文：{% for item in adjustments %}（{{ loop.index }}）{{ item.content }}，造价核减{{ item.amount }}万元；
            # 新文：{% for item in adjustments %}（{{ loop.index }}）{{ item.content }}，造价核减{{ item.amount }}万元{% if loop.last %}。{% else %}；{% endif %}

            if old_text.endswith('；'):
                new_text = old_text[:-1] + '{% if loop.last %}。{% else %}；{% endif %}'
            elif '万元；' in old_text:
                new_text = old_text.replace('万元；', '万元{% if loop.last %}。{% else %}；{% endif %}')
            else:
                print("⚠️  未找到预期的分号，请手动检查")
                return

            # 清空段落并重新设置文本
            para.clear()
            para.add_run(new_text)

            print(f"   新内容: {new_text}")
            break

    if not target_para:
        print("❌ 未找到包含 'adjustments' 的段落")
        return

    # 保存修改后的模板
    doc.save(template_path)
    print(f"✅ 模板标点符号已修复！")
    print(f"✅ 修改后的模板: {template_path}")
    print(f"💡 如需恢复，请使用备份文件: {backup_path}")

if __name__ == "__main__":
    try:
        fix_template_punctuation()
        print("\n✨ 修复完成！请重启服务并测试生成报告")
    except Exception as e:
        print(f"❌ 修复失败: {e}")
        import traceback
        traceback.print_exc()
