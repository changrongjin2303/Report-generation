#!/usr/bin/env python3
"""
将模板中的项目名称字段转换为表格格式
使用 python-docx 库自动修改 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def remove_table_borders(table):
    """移除表格所有边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # 创建无边框样式
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)

    tblPr.append(tblBorders)

def convert_template():
    """转换模板文件"""
    template_path = "/Users/crj/Documents/code/baogao_code/report_template.docx"
    backup_path = "/Users/crj/Documents/code/baogao_code/report_template_backup.docx"

    print("🔍 正在读取模板文件...")
    doc = Document(template_path)

    # 备份原模板
    doc.save(backup_path)
    print(f"✅ 已备份原模板到: {backup_path}")

    # 查找包含 "咨询项目全称" 的段落
    target_para = None
    for i, para in enumerate(doc.paragraphs):
        if "咨询项目全称" in para.text or "project_name" in para.text:
            target_para = para
            para_index = i
            print(f"✅ 找到目标段落（第{i+1}段）: {para.text[:50]}...")
            break

    if not target_para:
        print("❌ 未找到包含 '咨询项目全称' 的段落")
        return

    # 在目标段落之前插入表格
    print("📝 正在创建表格...")
    table = doc.add_table(rows=1, cols=2)

    # 移动表格到正确位置（在找到的段落之前）
    tbl_element = table._element
    target_para._element.addprevious(tbl_element)

    # 设置表格内容
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    left_cell.text = "咨询项目全称:"
    right_cell.text = "{{ project_name }}"

    # 设置左列宽度（4cm）
    left_cell.width = Cm(4)
    # 右列自动

    # 设置右列居中对齐
    right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置字体大小
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 移除表格边框
    remove_table_borders(table)

    # 删除原来的段落
    target_para._element.getparent().remove(target_para._element)

    # 保存修改后的模板
    doc.save(template_path)
    print(f"✅ 模板已成功转换为表格格式！")
    print(f"✅ 修改后的模板: {template_path}")
    print(f"💡 如需恢复，请使用备份文件: {backup_path}")

if __name__ == "__main__":
    try:
        convert_template()
        print("\n✨ 转换完成！请重启服务并测试生成报告")
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        import traceback
        traceback.print_exc()
